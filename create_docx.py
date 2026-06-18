#!/usr/bin/env python3
"""
Generate a professional Word document from report.md.

Features:
- Renders Markdown tables as real Word tables (no raw pipes).
- Embeds SVG charts as PNG images at relevant sections.
- Formats ASCII diagrams as styled code blocks (no backticks).
- Normalises German umlauts and cleans HTML entities.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
REPORT_MD = BASE_DIR / "report.md"
OUTPUT_DOCX = BASE_DIR / "AgroPark_Nekrasovo_Strategiepapier.docx"
PUBLIC_DIR = BASE_DIR / "public"
BUILD_DIR = BASE_DIR / ".docx_build"

BRAND_GREEN = RGBColor(0x1B, 0x43, 0x32)
BRAND_GREEN_LIGHT = RGBColor(0x2D, 0x6A, 0x4F)
BRAND_BROWN = RGBColor(0xD4, 0xA3, 0x73)
TEXT_DARK = RGBColor(0x2C, 0x2C, 0x2C)
TEXT_GRAY = RGBColor(0x5C, 0x5F, 0x62)
BG_LIGHT = "F5F5F0"

# Section heading -> SVG chart to insert after that heading.
CHART_INSERTIONS = {
    "4.1 Globaler Agritech-Markt": PUBLIC_DIR / "de_01_markt.svg",
    "4.2 Wettbewerbsanalyse: 2×2-Matrix": PUBLIC_DIR / "de_05_wettbewerb.svg",
    "3.2 Geschäftsmodell-Architektur": PUBLIC_DIR / "de_02_umsatz.svg",
    "6.2 Operativer Impact": PUBLIC_DIR / "de_03_ai_impact.svg",
    "7.2 Lokale SEO-Strategie": PUBLIC_DIR / "de_04_sichtbarkeit.svg",
}

# ---------------------------------------------------------------------------
# Font helpers
# ---------------------------------------------------------------------------
def set_run_font(run, name="Arial Narrow", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure Cyrillic / German characters render with the chosen font.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name="Arial Narrow", size=11, bold=False, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


# ---------------------------------------------------------------------------
# Text preprocessing
# ---------------------------------------------------------------------------
def normalize_umlauts(text):
    """Convert ae/oe/ue word parts to German umlauts, respecting word boundaries."""
    patterns = [
        (r"(?<![A-Za-z])ae(?![A-Za-z])", "ä"),
        (r"(?<![A-Za-z])oe(?![A-Za-z])", "ö"),
        (r"(?<![A-Za-z])ue(?![A-Za-z])", "ü"),
        (r"(?<![A-Za-z])Ae(?![A-Za-z])", "Ä"),
        (r"(?<![A-Za-z])Oe(?![A-Za-z])", "Ö"),
        (r"(?<![A-Za-z])Ue(?![A-Za-z])", "Ü"),
        # Common whole-word ss replacements where ß is expected.
        (r"(?<![A-Za-z])Strasse(?![A-Za-z])", "Straße"),
        (r"(?<![A-Za-z])strasse(?![A-Za-z])", "straße"),
        (r"(?<![A-Za-z])Massnahme(?![A-Za-z])", "Maßnahme"),
        (r"(?<![A-Za-z])massnahme(?![A-Za-z])", "maßnahme"),
        (r"(?<![A-Za-z])Ausschluss(?![A-Za-z])", "Ausschluß"),
        (r"(?<![A-Za-z])ausschluss(?![A-Za-z])", "ausschluß"),
    ]
    for pat, repl in patterns:
        text = re.sub(pat, repl, text)
    return text


def preprocess_line(text):
    """Clean entities and markdown shortcuts before rendering."""
    # HTML entities
    text = text.replace("&mdash;", "–")
    text = text.replace("&ndash;", "–")
    text = text.replace("&times;", "×")
    text = text.replace("&deg;", "°")
    text = text.replace("&amp;", "&")
    text = text.replace("&quot;", '"')
    text = text.replace("&#39;", "'")
    # Markdown links -> just the label
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Task list checkboxes
    text = text.replace("[ ]", "☐")
    text = text.replace("[x]", "☑")
    # Umlauts
    text = normalize_umlauts(text)
    return text


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------
def add_inline_text(paragraph, text):
    """Add text to a paragraph, handling **bold**, *italic* and `code`."""
    # Escape asterisks that are not paired? We'll use a simple splitter.
    # Process segments in order: bold, italic, code.
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Courier New", size=10)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------
def add_normal_paragraph(doc, text, style="Normal", alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    add_inline_text(p, text)
    return p


def add_heading(doc, text, level):
    style_name = f"Heading {level}"
    p = doc.add_paragraph(style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    add_inline_text(p, text)
    return p


def add_blockquote(doc, lines):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)
    # Add a left border via paragraph shading? Use a table instead for a clean look.
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, BG_LIGHT)
    set_cell_borders(cell, color="D4A373", sz=8)
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
    full_text = " ".join(lines)
    run = p_cell.add_run(full_text)
    set_run_font(run, italic=True, color=TEXT_DARK)
    # Remove empty paragraph after table
    return table


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(sz))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table, color="CCCCCC", sz=4):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color=color, sz=sz)


def parse_table_row(line):
    """Split a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().split("|")]
    # Remove empty cells caused by leading/trailing pipes.
    cells = [c for c in cells if c != "" or c == ""]
    # Trim exactly one empty cell at each end if present.
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def is_separator_row(cells):
    """True if every cell consists of dashes and optional colons only."""
    if not cells:
        return False
    for c in cells:
        if not re.fullmatch(r":?-+:?", c.strip()):
            return False
    return True


def add_markdown_table(doc, raw_rows):
    """raw_rows: list of original markdown table lines."""
    parsed_rows = [parse_table_row(line) for line in raw_rows]
    if not parsed_rows:
        return

    # Detect header separator row.
    header_idx = None
    for i, row in enumerate(parsed_rows):
        if is_separator_row(row):
            header_idx = i
            break

    if header_idx is not None:
        header_rows = parsed_rows[:header_idx]
        body_rows = parsed_rows[header_idx + 1:]
    else:
        header_rows = []
        body_rows = parsed_rows

    # If there's an empty header row, skip it.
    if header_rows and all(c == "" for c in header_rows[0]):
        header_rows = header_rows[1:]

    # Determine column count.
    col_count = max(len(r) for r in parsed_rows) if parsed_rows else 1

    # Build data for the Word table.
    word_rows = []
    if header_rows:
        word_rows.append(pad_row(header_rows[0], col_count))
    for r in body_rows:
        word_rows.append(pad_row(r, col_count))

    if not word_rows:
        return

    table = doc.add_table(rows=len(word_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table, color="D9D9D6", sz=4)

    # Column width: fit within 6.3 inches.
    page_width = Inches(6.3)
    col_width = page_width / col_count

    for i, row in enumerate(word_rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.width = col_width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if i == 0 and header_rows:
                set_cell_shading(cell, "1B4332")
                run = p.add_run(cell_text)
                set_run_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            else:
                # Alternate row shading for readability.
                if i % 2 == 0:
                    set_cell_shading(cell, "FAFAF8")
                add_inline_text(p, cell_text)

    # Add spacing after table.
    doc.add_paragraph()
    return table


def pad_row(row, n):
    return row + [""] * (n - len(row))


# ---------------------------------------------------------------------------
# Code block / diagram helpers
# ---------------------------------------------------------------------------
def add_code_block(doc, lines, caption=None):
    """Render a code block as a shaded, bordered monospace box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    set_cell_shading(cell, BG_LIGHT)
    set_cell_borders(cell, color="1B4332", sz=6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line.rstrip())
        set_run_font(run, name="Courier New", size=9, color=TEXT_DARK)

    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(4)
        cap_p.paragraph_format.space_after = Pt(12)
        run = cap_p.add_run(caption)
        set_run_font(run, size=10, italic=True, color=TEXT_GRAY)
    else:
        doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------
def ensure_png(svg_path):
    """Convert SVG to PNG (cached in .docx_build) and return PNG path."""
    if not svg_path.exists():
        return None
    BUILD_DIR.mkdir(exist_ok=True)
    png_path = BUILD_DIR / (svg_path.stem + ".png")
    if png_path.exists() and png_path.stat().st_mtime >= svg_path.stat().st_mtime:
        return png_path
    try:
        import cairosvg
        cairosvg.svg2png(url=str(svg_path), write_to=str(png_path), output_width=1600)
    except Exception as exc:
        print(f"Warning: could not convert {svg_path}: {exc}")
        return None
    return png_path


def add_chart(doc, svg_path, caption):
    png_path = ensure_png(svg_path)
    if not png_path:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(5.8))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(14)
    run = cap_p.add_run(caption)
    set_run_font(run, size=10, italic=True, color=TEXT_GRAY)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_document_styles(doc):
    # Normal
    set_style_font(doc.styles["Normal"], name="Arial Narrow", size=11, color=TEXT_DARK)
    doc.styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    doc.styles["Normal"].paragraph_format.space_after = Pt(8)

    # Heading styles
    for level, size, color in [(1, 22, BRAND_GREEN), (2, 16, BRAND_GREEN), (3, 13, BRAND_GREEN_LIGHT)]:
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, name="Arial Narrow", size=size, bold=True, color=color)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    # List styles
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, name="Arial Narrow", size=11, color=TEXT_DARK)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(4)


def build_document():
    doc = Document()
    setup_document_styles(doc)

    # Page margins
    sections = doc.sections[0]
    sections.top_margin = Inches(0.9)
    sections.bottom_margin = Inches(0.9)
    sections.left_margin = Inches(0.9)
    sections.right_margin = Inches(0.9)

    with open(REPORT_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    current_heading = ""
    i = 0
    n = len(lines)

    while i < n:
        raw_line = lines[i]
        line = raw_line.rstrip("\n")
        stripped = line.strip()
        proc = preprocess_line(stripped)

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"-+", stripped):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = preprocess_line(m.group(2))
            add_heading(doc, text, level)
            current_heading = text
            # Insert chart if this heading triggers one.
            for key, svg_path in CHART_INSERTIONS.items():
                if key in current_heading or current_heading in key:
                    add_chart(doc, svg_path, caption=f"Abbildung: {text}")
                    break
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(preprocess_line(lines[i].rstrip("\n")))
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        # Code blocks
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(preprocess_line(lines[i].rstrip("\n")))
                i += 1
            i += 1  # skip closing fence
            caption = f"Abbildung: {current_heading}" if "Diagramm" in current_heading else None
            add_code_block(doc, code_lines, caption=caption)
            continue

        # Blockquotes
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(preprocess_line(lines[i].strip().lstrip(">").strip()))
                i += 1
            add_blockquote(doc, quote_lines)
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            text = preprocess_line(m_num.group(2))
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_inline_text(p, text)
            i += 1
            continue

        # Bullet list
        if stripped.startswith(("- ", "* ")):
            text = preprocess_line(stripped[2:])
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline_text(p, text)
            i += 1
            continue

        # Regular paragraph
        paragraphs = [proc]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(("#", "|", "```", ">", "- ", "* ")) and not re.match(r"^\d+\.", lines[i].strip()):
            paragraphs.append(preprocess_line(lines[i].strip()))
            i += 1
        add_normal_paragraph(doc, " ".join(paragraphs))

    doc.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_document()
